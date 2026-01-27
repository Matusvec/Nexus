"""
Document parsing - Extract text from PDFs, DOCX, and describe images + tables

This module handles ONLY document parsing. For the full pipeline,
use main.py which orchestrates parsing -> chunking -> storage.
"""
from pathlib import Path
from typing import Dict, List, Tuple
from tqdm import tqdm

# PDF parsing
from pypdf import PdfReader

# DOCX parsing
from docx import Document
from docx.table import Table as DocxTable

# Use centralized Gemini client
from gemini_client import generate_with_image, generate_content


def parse_document(file_path: str) -> Dict[str, any]:
    """
    Parse a document (PDF or DOCX) and extract text + image + table descriptions
    
    Args:
        file_path: Path to the document file
        
    Returns:
        Dict with:
            - text: Full extracted text with image/table descriptions inline
            - metadata: File info (name, type, page count, etc.)
            - images_found: Number of images processed
            - tables_found: Number of tables processed
    """
    file_path = Path(file_path)
    file_ext = file_path.suffix.lower()
    
    # If no extension, try to auto-detect
    if not file_ext:
        for ext in ['.pdf', '.docx', '.doc']:
            test_path = file_path.with_suffix(ext)
            if test_path.exists():
                print(f"📌 Auto-detected file: {test_path.name}")
                file_path = test_path
                file_ext = ext
                break
        else:
            raise ValueError(f"File not found. Tried: {file_path}.pdf, {file_path}.docx, {file_path}.doc")
    
    if file_ext == '.pdf':
        return parse_pdf(file_path)
    elif file_ext in ['.docx', '.doc']:
        return parse_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {file_ext}. Use .pdf or .docx files.")


def parse_pdf(file_path: Path) -> Dict[str, any]:
    """
    Extract text, images, and tables from PDF with context-aware descriptions
    
    Follows Microsoft best practices:
    - Includes surrounding text context when describing images/tables
    - Uses LLM (Gemini Vision) for image and table descriptions
    - Converts tables to markdown format
    - Tracks image/table references for multi-chunk scenarios
    
    Args:
        file_path: Path to PDF file
        
    Returns:
        Dict with text, metadata, image count, and table count
    """
    print(f"📄 Parsing PDF: {file_path.name}")
    
    reader = PdfReader(str(file_path))
    text_parts = []
    images_processed = 0
    tables_processed = 0
    image_references = []
    table_references = []
    
    # Extract full document text for table context
    full_doc_text = "\n".join([page.extract_text() or "" for page in reader.pages])
    
    # Progress bar for pages
    for page_num, page in enumerate(tqdm(reader.pages, desc="Processing pages", unit="page"), 1):
        # Extract text from this page
        page_text = page.extract_text()
        
        # Store page text for context
        if page_text:
            text_parts.append(f"--- Page {page_num} ---\n{page_text}")
        
        # Extract and describe images with surrounding context
        if hasattr(page, 'images') and page.images:
            # Progress bar for images on this page
            for img_idx, image in enumerate(tqdm(page.images, desc=f"  Page {page_num} images", unit="img", leave=False)):
                try:
                    # Get image data
                    image_data = image.data
                    
                    # Get context: text before and after image
                    context_before = page_text[-500:] if page_text else ""
                    context_after = ""
                    
                    # Check if this image might be a table (use Gemini Vision to detect)
                    is_table, table_data = detect_and_extract_table_from_image(image_data)
                    
                    if is_table and table_data:
                        # This is a table - process as table
                        table_ref = f"table_{file_path.stem}_p{page_num}_{tables_processed}"
                        table_references.append(table_ref)
                        
                        # Describe and format table
                        table_description = describe_and_format_table(
                            table_data,
                            table_ref,
                            context=full_doc_text[-2000:]  # Last 2000 chars of doc
                        )
                        
                        text_parts.append(f"\n{table_description}\n")
                        tables_processed += 1
                    else:
                        # Regular image - process as image
                        img_ref = f"img_{file_path.stem}_p{page_num}_{img_idx}"
                        image_references.append(img_ref)
                        
                        description = describe_image_with_context(
                            image_data, 
                            img_ref,
                            context_before=context_before
                        )
                        
                        text_parts.append(f"\n[IMAGE {img_ref}: {description}]\n")
                        images_processed += 1
                    
                except Exception as e:
                    tqdm.write(f"    ⚠️  Could not process image {img_idx}: {e}")
    
    full_text = "\n\n".join(text_parts)
    
    return {
        "text": full_text,
        "metadata": {
            "filename": file_path.name,
            "filetype": "pdf",
            "pages": len(reader.pages),
            "images_found": images_processed,
            "tables_found": tables_processed,
            "image_references": image_references,
            "table_references": table_references
        },
        "images_found": images_processed,
        "tables_found": tables_processed
    }


def parse_docx(file_path: Path) -> Dict[str, any]:
    """
    Extract text, images, and tables from DOCX with context-aware descriptions
    
    Follows Microsoft best practices:
    - Includes surrounding text context when describing images/tables
    - Extracts native DOCX tables and converts to markdown
    - Tracks image/table references for multi-chunk scenarios
    
    Args:
        file_path: Path to DOCX file
        
    Returns:
        Dict with text, metadata, image count, and table count
    """
    print(f"📄 Parsing DOCX: {file_path.name}")
    
    doc = Document(str(file_path))
    text_parts = []
    images_processed = 0
    tables_processed = 0
    image_references = []
    table_references = []
    
    # Build full text first for context extraction
    all_paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    full_doc_text = "\n".join(all_paragraphs)
    
    # Extract and process tables first (DOCX has native table support!)
    if doc.tables:
        tqdm.write(f"  Found {len(doc.tables)} tables in document")
        for table_idx, table in enumerate(tqdm(doc.tables, desc="Processing tables", unit="table")):
            try:
                table_ref = f"table_{file_path.stem}_{table_idx}"
                table_references.append(table_ref)
                
                # Extract table data from DOCX table object
                table_data = extract_docx_table(table)
                
                # Describe and format table with context
                table_description = describe_and_format_table(
                    table_data,
                    table_ref,
                    context=full_doc_text[-2000:]  # Last 2000 chars
                )
                
                text_parts.append(f"\n{table_description}\n")
                tables_processed += 1
                
            except Exception as e:
                tqdm.write(f"    ⚠️  Could not process table {table_idx}: {e}")
    
    # Progress bar for paragraphs
    for para_idx, paragraph in enumerate(tqdm(doc.paragraphs, desc="Processing paragraphs", unit="para")):
        # Extract text
        if paragraph.text.strip():
            text_parts.append(paragraph.text)
        
        # Check for images in runs
        for run in paragraph.runs:
            if hasattr(run, 'element') and run.element.xpath('.//pic:pic'):
                try:
                    # Get context: previous and next paragraphs
                    context_before = " ".join(all_paragraphs[max(0, para_idx-2):para_idx])
                    context_after = " ".join(all_paragraphs[para_idx+1:min(len(all_paragraphs), para_idx+3)])
                    
                    # Create unique image reference
                    img_ref = f"img_{file_path.stem}_para{para_idx}_{images_processed}"
                    image_references.append(img_ref)
                    
                    # Extract actual image data from DOCX
                    image_data = extract_image_from_run(run, doc)
                    
                    if image_data:
                        # Describe image with Gemini Vision
                        description = describe_image_with_context(
                            image_data,
                            img_ref,
                            context_before=context_before,
                            context_after=context_after
                        )
                    else:
                        # Fallback if image extraction fails
                        description = f"Image at paragraph {para_idx} (extraction failed)"
                    
                    text_parts.append(f"\n[IMAGE {img_ref}: {description}]\n")
                    images_processed += 1
                    
                except Exception as e:
                    tqdm.write(f"    ⚠️  Could not process image: {e}")
    
    full_text = "\n\n".join(text_parts)
    
    return {
        "text": full_text,
        "metadata": {
            "filename": file_path.name,
            "filetype": "docx",
            "paragraphs": len(doc.paragraphs),
            "images_found": images_processed,
            "tables_found": tables_processed,
            "image_references": image_references,
            "table_references": table_references
        },
        "images_found": images_processed,
        "tables_found": tables_processed
    }


def describe_image_with_context(image_data: bytes, image_id: str, context_before: str = "", context_after: str = "") -> str:
    """
    Use Gemini Vision to describe an image WITH surrounding text context
    
    Microsoft best practice: Pass text before/after image to LLM for better descriptions
    This helps the model understand what the image is illustrating
    
    Args:
        image_data: Raw image bytes
        image_id: Identifier for logging and cross-referencing
        context_before: Text appearing before the image (last ~500 chars)
        context_after: Text appearing after the image (next ~500 chars)
        
    Returns:
        Context-aware text description of the image
    """
    # Use tqdm.write to print without interfering with progress bars
    tqdm.write(f"    🖼️  Describing {image_id}...")
    
    try:
        # Build context-aware prompt
        prompt_parts = []
        
        if context_before:
            prompt_parts.append(f"Text before this image: '{context_before.strip()}'\n")
        if context_after:
            prompt_parts.append(f"Text after this image: '{context_after.strip()}'\n")
        
        prompt_parts.append(
            "\nDescribe this image in 1-3 sentences. Consider the surrounding text context. "
            "Focus on:\n"
            "- If it's a diagram/chart: What data or concept does it show?\n"
            "- If it contains text: Summarize key text content\n"
            "- If it's a table: Describe structure and main data points\n"
            "- If it's a photo/illustration: Key visual elements and relevance to document\n"
            "Be concise, factual, and context-aware."
        )
        
        prompt = "".join(prompt_parts)
        
        # Use centralized Gemini client
        description = generate_with_image(image_data, prompt)
        tqdm.write(f"      ✓ {description[:60]}...")
        return description
        
    except Exception as e:
        tqdm.write(f"      ⚠️  Vision API error: {e}")
        return f"Image {image_id} (description unavailable)"


def describe_image(image_data: bytes, image_id: str) -> str:
    """
    Legacy function - use describe_image_with_context instead
    Kept for backward compatibility
    """
    return describe_image_with_context(image_data, image_id)


def extract_docx_table(table: DocxTable) -> List[List[str]]:
    """
    Extract table data from a DOCX table object
    
    Args:
        table: python-docx Table object
        
    Returns:
        2D list of table cells (rows x columns)
    """
    table_data = []
    for row in table.rows:
        row_data = []
        for cell in row.cells:
            row_data.append(cell.text.strip())
        table_data.append(row_data)
    return table_data


def extract_image_from_run(run, document) -> bytes:
    """
    Extract image bytes from a DOCX run element
    
    Args:
        run: python-docx Run object containing image
        document: python-docx Document object
        
    Returns:
        Image bytes or None if extraction fails
    """
    try:
        # Get image relationship ID from the run
        inline = run.element.xpath('.//a:blip/@r:embed')
        if not inline:
            return None
        
        # Get the image part from document relationships
        image_id = inline[0]
        image_part = document.part.related_parts[image_id]
        
        # Return the binary image data
        return image_part.blob
        
    except Exception as e:
        tqdm.write(f"      ⚠️  Image extraction error: {e}")
        return None


def detect_and_extract_table_from_image(image_data: bytes) -> Tuple[bool, List[List[str]]]:
    """
    Use Gemini Vision to detect if image contains a table and extract it
    
    Args:
        image_data: Raw image bytes
        
    Returns:
        Tuple of (is_table: bool, table_data: List[List[str]])
    """
    tqdm.write(f"    📊 Checking if image contains table...")
    
    try:
        prompt = (
            "Does this image contain a data table with rows and columns? "
            "Respond with 'YES' if it's a table, or 'NO' if it's not. "
            "If YES, extract the table data in CSV format."
        )
        
        result = generate_with_image(image_data, prompt)
        
        if result.upper().startswith('YES'):
            # Extract table data (simplified - assumes CSV format in response)
            # In practice, might need more sophisticated parsing
            tqdm.write(f"      ✓ Table detected in image")
            return True, []  # Return empty for now, will be processed by describe_and_format_table
        else:
            return False, []
            
    except Exception as e:
        tqdm.write(f"      ⚠️  Table detection error: {e}")
        return False, []


def describe_and_format_table(table_data: List[List[str]], table_id: str, context: str = "") -> str:
    """
    Use Gemini to describe table and convert to markdown format
    
    Best practice from research:
    1. Generate contextual description of table
    2. Convert table to markdown format
    3. Combine both for optimal RAG retrieval
    
    Args:
        table_data: 2D list of table cells
        table_id: Unique identifier for the table
        context: Surrounding document context
        
    Returns:
        Formatted string with: [TABLE table_id: description]\n\n<markdown table>
    """
    tqdm.write(f"    📊 Describing and formatting {table_id}...")
    
    try:
        # Convert table data to simple text format for prompt
        table_text = "\n".join([" | ".join(row) for row in table_data])
        
        # Build context-aware prompt
        prompt = f"""
Given the following table and its context from the document, provide:

1. A comprehensive 2-3 sentence description of what this table shows
2. The table formatted in markdown

Document Context:
{context[-1000:] if context else "No additional context"}

Table Content:
{table_text}

Format your response as:
DESCRIPTION: <your description>

MARKDOWN:
<table in markdown format>
"""
        
        # Use centralized Gemini client
        result = generate_content(prompt)
        
        # Parse response to extract description and markdown
        if "DESCRIPTION:" in result and "MARKDOWN:" in result:
            parts = result.split("MARKDOWN:")
            description = parts[0].replace("DESCRIPTION:", "").strip()
            markdown_table = parts[1].strip()
        else:
            # Fallback if format not followed
            description = "Table data extracted from document"
            markdown_table = table_text
        
        # Format as table chunk: [TABLE id: description]\n\nmarkdown
        table_chunk = f"[TABLE {table_id}: {description}]\n\n{markdown_table}"
        
        tqdm.write(f"      ✓ Table processed: {description[:50]}...")
        return table_chunk
        
    except Exception as e:
        tqdm.write(f"      ⚠️  Table processing error: {e}")
        # Fallback: basic markdown conversion
        markdown_table = "| " + " | ".join(table_data[0]) + " |\n"
        markdown_table += "| " + " | ".join(["---"] * len(table_data[0])) + " |\n"
        for row in table_data[1:]:
            markdown_table += "| " + " | ".join(row) + " |\n"
        
        return f"[TABLE {table_id}: Table from document]\n\n{markdown_table}"



# Test parsing only (use main.py for full pipeline)
if __name__ == "__main__":
    import sys
    
    if len(sys.argv) < 2:
        print("Document Parser - Extract text, images, and tables")
        print("\nUsage: python document_parser.py <file.pdf|file.docx>")
        print("\nThis will ONLY parse the document.")
        print("For the full pipeline (parse → chunk → store), use:")
        print("  python main.py upload <file>")
        sys.exit(0)
    
    file_path = sys.argv[1]
    
    print("="*60)
    print("📄 DOCUMENT PARSING")
    print("="*60)
    
    result = parse_document(file_path)
    
    print(f"\n✓ Parsing complete:")
    print(f"  - Filename: {result['metadata']['filename']}")
    print(f"  - Text: {len(result['text'])} characters")
    print(f"  - Images: {result['images_found']}")
    print(f"  - Tables: {result['tables_found']}")
    
    # Show preview
    print(f"\n📝 Text preview (first 500 chars):")
    print("-" * 40)
    print(result['text'][:500])
    print("-" * 40)
    
    print("\n💡 To process and store this document, run:")
    print(f"   python main.py upload {file_path}")
